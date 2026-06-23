import json
import logging
import os
import re
import shutil
import sys
from datetime import date
from typing import Any

import requests
import textstat
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s  %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger(__name__)

# Paths & Configurations 
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))


def _load(path: str) -> Any:
    with open(os.path.join(ROOT, path)) as f:
        return json.load(f)


SETTINGS    = _load("config/settings.json")
SCHEMA      = _load("config/template_schema.json")
BOILERPLATE = _load("config/boilerplate.json")
T2          = SETTINGS["track2"]
T3          = SETTINGS["track3"]
PATHS       = SETTINGS["paths"]

# Track 2 config
LLM_MODEL    = T2["llm_model"]
LLM_BASE_URL = T2["llm_base_url"]
OUTPUT_PATH  = os.path.join(ROOT, PATHS["proposal_output"])

# Track 3 config
T3_LLM_MODEL       = T3["llm_model"]
T3_LLM_BASE_URL    = T3["llm_base_url"]
READABILITY_MIN     = T3["readability_min_score"]
DRAFT_PATH          = os.path.join(ROOT, PATHS["proposal_output"])
VALIDATED_PATH      = os.path.join(ROOT, PATHS["validated_output"])
COMPLETENESS_REPORT = os.path.join(ROOT, PATHS["completeness_report"])
CONSISTENCY_REPORT  = os.path.join(ROOT, PATHS["consistency_report"])
READABILITY_REPORT  = os.path.join(ROOT, PATHS["readability_report"])


# llm caller
def call_llm(prompt: str, model: str = None, base_url: str = None) -> str:

    _model   = model    or LLM_MODEL
    _baseurl = base_url or LLM_BASE_URL

    try:
        resp = requests.post(
            f"{_baseurl}/api/generate",
            json={"model": _model, "prompt": prompt, "stream": False},
            timeout=300,
        )
        resp.raise_for_status()
        return resp.json()["response"].strip()

    except requests.exceptions.ConnectionError:
        log.warning("Ollama not running. Using placeholder text.")
        return "[LLM OUTPUT — run Ollama with: ollama run %s]" % _model

    except Exception as e:
        log.error("LLM call failed: %s", e)
        return f"[LLM ERROR: {e}]"


# load inputs
def load_inputs(team1_path: str, team2_path: str) -> tuple[dict, dict]:
# in reality the input would be from a DB or API . Change acc
    log.info("Loading Team 1 scope model: %s", team1_path)
    team1 = _load(team1_path)

    log.info("Loading Team 2 cost estimate: %s", team2_path)
    team2 = _load(team2_path)

    return team1, team2


# Aggregate fn : Returns a context dict (ctx) that LLM will reference for any gen.
    
def aggregate_scope(team1: dict, team2: dict) -> dict:
    log.info("Aggregating scope and cost data...")


    by_material: dict[str, dict] = {}
    for item in team1.get("scope_items", []):
        mat = item["material"]
        if mat not in by_material:
            by_material[mat] = {"areas": [], "total_qty": 0, "unit": item["unit"]}
        by_material[mat]["areas"].append(item["area"])
        by_material[mat]["total_qty"] += item["quantity"]


    scope_lines = []
    for mat, data in by_material.items():
        areas = ", ".join(data["areas"])
        scope_lines.append(
            f"{data['total_qty']:,} {data['unit']} of {mat} in {areas}"
        )


    pricing_lines = []
    for row in team2.get("cost_estimate", []):
        pricing_lines.append(
            f"  Item {row['item']}: {row['description']} — "
            f"{row['quantity']} @ {row['unit_cost']} = {row['total']}"
        )

    ctx = {
        "project":        team1["project"],
        "scope_summary":  "\n".join(scope_lines),
        "pricing_lines":  "\n".join(pricing_lines),
        "base_bid_total": team2["base_bid_total"],
        "scope_gaps":     team1.get("scope_gaps", []),
        "conflicts":      team2.get("conflicts", []),
        "risk_flags":     team2.get("risk_flags", []),
        "alternates":     team1.get("alternates", []),
        "cost_rows":      team2.get("cost_estimate", []),
    }

    log.info("Aggregated %d scope items, base bid total: %s",
             len(team1.get("scope_items", [])), team2["base_bid_total"])
    return ctx


# LLM generation - prompts
PROMPTS = {
    "scope_of_work": """You are a professional construction estimator writing a flooring proposal.
Write a concise technical Scope of Work section (3-5 bullet points) for this project.
Use professional contractor language. Be specific about materials, areas and quantities.

Project: {project_name}
Client: {client}
Scope items:
{scope_summary}

Output bullet points only. No headers. No intro sentence.""",

    "inclusions": """You are a professional construction estimator.
Write an Inclusions section that explicitly states what IS covered beyond
the line items — things the client might otherwise assume are extra.
Do NOT just repeat the scope items.

Include things like: surface protection after installation, final cleanup,
transition strips, warranty coverage, primer and adhesive, substrate prep.

Project scope for context:
{scope_summary}

5-7 bullet points. Output bullet points only.""",

    "exclusions": """You are a professional construction estimator.
Write an Exclusions section listing what is explicitly NOT included.
Be specific — name the actual conditions or areas being excluded.
Vague exclusions like "moisture mitigation for unlisted areas" are not acceptable.

Scope gaps identified:
{scope_gaps}

Conflicts flagged:
{conflicts}

Also exclude: furniture removal, overtime work, structural substrate repairs,
permits unless specified, work outside normal business hours.

5-7 bullet points. Output bullet points only.""",

    "qualifications": """You are a professional construction estimator.
Write a Qualifications section listing the site conditions and assumptions
under which this proposal price is valid. These are contractual preconditions,
NOT company credentials or resume content.

Examples of the correct style:
- Concrete substrates assumed flat within 3/16 inch over 10 ft radius.
- HVAC must be operational maintaining 65-85°F during and 48 hours post-installation.
- Work areas must be free of other trades during flooring installation.

Risk flags to incorporate:
{risk_flags}

Project: {project_name}, Location: {location}

4-6 bullet points. Output bullet points only.""",

    "alternate_pricing": """You are a professional construction estimator.
Write a brief description for each alternate pricing option listed below.
One sentence per alternate. Professional tone.

Alternates:
{alternates}

Output one line per alternate in format: Alt-X: [description] — [delta]"""
}

# LLM generation - section
def generate_section(section_id: str, ctx: dict) -> str:

    if section_id not in PROMPTS:
        log.warning("No prompt defined for section: %s", section_id)
        return ""

    p = ctx["project"]
    prompt = PROMPTS[section_id].format(
        project_name  = p["name"],
        client        = p["client"],
        location      = p["location"],
        scope_summary = ctx["scope_summary"],
        scope_gaps    = "\n".join(f"- {g}" for g in ctx["scope_gaps"]),
        conflicts     = "\n".join(f"- {c}" for c in ctx["conflicts"]),
        risk_flags    = "\n".join(f"- {r}" for r in ctx["risk_flags"]),
        alternates    = "\n".join(
            f"- {a['id']}: {a['description']} ({a['delta']})"
            for a in ctx["alternates"]
        ),
    )

    log.info("Generating section: %s", section_id)
    text = call_llm(prompt)
    log.info("Section %s: %d chars generated", section_id, len(text))
    return text


# make the doc file in memory
def _heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.size = Pt(16 if level == 1 else 13)


def _bullet(doc: Document, text: str):
    clean = text.lstrip("-•* ").strip()
    if not clean:
        return
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(clean).font.size = Pt(11)


def _para(doc: Document, text: str):
    p = doc.add_paragraph()
    p.add_run(text).font.size = Pt(11)


def _add_header(doc: Document, project: dict):
    doc.add_heading("FLOORING INSTALLATION PROPOSAL", 0)
    doc.add_paragraph()

    fields = [
        ("Prepared By",      project["contractor"]),
        ("Prepared For",     project["client"]),
        ("Project",          project["name"]),
        ("Bid Number",       project["bid_number"]),
        ("Date Submitted",   project["submission_date"]),
        ("Project Location", project["location"]),
        ("Scope Type",       project["scope_type"]),
    ]
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(fields):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()


def _add_pricing_table(doc: Document, cost_rows: list, total: str):
    headers = ["#", "Description", "Quantity", "Unit Cost", "Total"]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    for row in cost_rows:
        cells = table.add_row().cells
        cells[0].text = str(row["item"])
        cells[1].text = row["description"]
        cells[2].text = row["quantity"]
        cells[3].text = row["unit_cost"]
        cells[4].text = row["total"]

    # Total row
    tr = table.add_row().cells
    tr[0].text = ""
    tr[1].text = "BASE BID TOTAL"
    tr[1].paragraphs[0].runs[0].bold = True
    tr[2].text = ""
    tr[3].text = ""
    tr[4].text = total
    tr[4].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()


def _add_llm_section(doc: Document, label: str, text: str):
    _heading(doc, label)
    for line in text.split("\n"):
        line = line.strip()
        if line:
            _bullet(doc, line)
    doc.add_paragraph()


def _add_terms(doc: Document, boilerplate: dict):
    # _heading(doc, "Terms and Conditions")
    for term in boilerplate["terms"]:
        _bullet(doc, term)
    doc.add_paragraph()
    _para(doc, boilerplate["closing"])
    doc.add_paragraph()
    sig = boilerplate["signature_block"]
    _para(doc, f"{sig['name']}")
    _para(doc, f"{sig['company']}")
    _para(doc, f"{sig['phone']}")


def assemble_document(ctx: dict, generated: dict) -> Document:
    
    log.info("Assembling proposal document...")
    doc = Document()

    # Set default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    section_num = 1
    for section in SCHEMA["sections"]:
        sid     = section["id"]
        label   = section["label"]
        use_llm = section["llm"]

        if sid == "header":
            _add_header(doc, ctx["project"])

        elif sid == "pricing_summary":
            _heading(doc, f"{section_num}. {label}")
            _add_pricing_table(doc, ctx["cost_rows"], ctx["base_bid_total"])
            section_num += 1

        elif sid == "terms":
            _heading(doc, f"{section_num}. {label}")
            _add_terms(doc, BOILERPLATE)
            section_num += 1

        elif use_llm and sid in generated:
            _heading(doc, f"{section_num}. {label}")
            _add_llm_section(doc, "", generated[sid])  # heading already added
            section_num += 1

        elif not use_llm:
            log.debug("Skipping non-LLM section with no special handler: %s", sid)

    log.info("Document assembled: %d sections", section_num - 1)
    return doc

# save to disk
def save_document(doc: Document, path: str):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    log.info("Proposal saved → %s", path)


# load draft - generated_prposal.docx for validation 
def load_draft(path: str) -> dict[str, str]:

    log.info("Loading draft proposal: %s", path)
    doc = Document(path)

    sections: dict[str, str] = {}
    current_heading = None
    buffer = []

    for para in doc.paragraphs:
        style = para.style.name
        text  = para.text.strip()
        if not text:
            continue

        if "Heading" in style or style == "Title":
            # Save previous section
            if current_heading:
                sections[current_heading] = "\n".join(buffer).strip()
            current_heading = re.sub(r"^\d+\.\s*", "", text)
            buffer = []
        else:
            buffer.append(text)

    # Last section
    if current_heading:
        sections[current_heading] = "\n".join(buffer).strip()

    log.info("Extracted %d sections from draft", len(sections))
    return sections

# for future DB integration
def load_team2(team2_path: str) -> dict:
    log.info("Loading Team 2 cost data: %s", team2_path)
    return _load(team2_path)


# completeness check
def check_completeness(sections: dict[str, str]) -> list[str]:

    log.info("Running completeness check...")
    issues = []

    required = [s for s in SCHEMA["sections"] if s.get("required", True)]

    for s in required:
        label = s["label"]
        sid   = s["id"]

        # match sections
        matched = next(
            (k for k in sections if label.lower() in k.lower()),
            None
        )

        if not matched:
            issues.append(f"MISSING section: '{label}' (id: {sid})")
        elif not sections[matched].strip():
            issues.append(f"EMPTY section: '{label}' (id: {sid})")

    if issues:
        log.warning("Completeness issues found: %d", len(issues))
    else:
        log.info("Completeness check passed.")

    return issues

# consistency check - 
def _extract_currency_amounts(text: str) -> list[str]:
    return re.findall(r"(?:Rs\.\s*|\$)[\d,]+(?:\.\d{1,2})?",text)


def check_consistency(sections: dict[str, str], team2: dict) -> list[str]:

    log.info("Running consistency check...")
    issues = []

    # Collect all valid amounts from Team 2
    valid_amounts: set[str] = set()
    for row in team2.get("cost_estimate", []):
        # valid_amounts.add(float(row["total"])) # for number cost in input data.
        valid_amounts.add(row["total"].replace(" ", ""))
        valid_amounts.add(row["unit_cost"].replace(" ", ""))
    valid_amounts.add(team2.get("base_bid_total", "").replace(" ", ""))

    for alt in team2.get("alternates", []):
        if "delta" in alt:
            valid_amounts.add(alt["delta"].replace(" ", ""))

    for label, text in sections.items():
        if label in ("Pricing Summary", "Terms and Conditions", "FLOORING INSTALLATION PROPOSAL"):
            continue  # skip table sections and header
        found = _extract_currency_amounts(text)
        for amt in found:
            normalized = amt.replace(" ", "")
            if normalized not in valid_amounts:
                issues.append(
                    f"Unverified amount {amt} in section '{label}' "
                    f"— not found in Team 2 cost data."
                )

    # Check base bid total is present somewhere in pricing section
    pricing_text = next(
        (v for k, v in sections.items() if "pricing" in k.lower()), ""
    )
    expected_total = team2.get("base_bid_total", "")
    if expected_total and expected_total not in pricing_text:
        issues.append(
            f"Base bid total {expected_total} not found in Pricing Summary section."
        )

    if issues:
        log.warning("Consistency issues found: %d", len(issues))
    else:
        log.info("Consistency check passed.")

    return issues


# contradiction check - LLM
CONTRADICTION_PROMPT = """You are a senior proposal reviewer at a construction estimating firm.
Read the following proposal draft carefully and identify any logical contradictions between sections.

Examples of contradictions to look for:
- An item excluded in the Exclusions section but priced in Pricing Summary
- A material mentioned in Scope of Work but missing from Inclusions
- A qualification that conflicts with the Terms and Conditions

Proposal draft:
{draft_text}

List only genuine contradictions. Format each as:
- [Section A] vs [Section B]: <description of contradiction>

If no contradictions are found, output exactly: NO CONTRADICTIONS FOUND
This is a judgment call — flag only clear logical conflicts, not style issues."""


def check_contradictions(sections: dict[str, str]) -> str:

    log.info("Running LLM contradiction check...")

    draft_text = "\n\n".join(
        f"=== {label} ===\n{text}"
        for label, text in sections.items()
        if text.strip()
    )

    prompt = CONTRADICTION_PROMPT.format(draft_text=draft_text)
    result = call_llm(prompt, model=T3_LLM_MODEL, base_url=T3_LLM_BASE_URL)
    log.info("Contradiction check complete.")
    return result

# readability check + LLM to suggest rewrites
#textstat Flesch Reading Ease : objective
#call_llm : subjective
REWRITE_PROMPT = """You are a professional proposal editor.
The following section of a construction proposal scored poorly on readability.
Rewrite it to be clearer and more concise while keeping all technical content accurate.
Do not change any numbers, material names, or quantities.

Section: {section_label}
Original text:
{section_text}

Output the rewritten version only. No explanation."""


def check_readability(sections: dict[str, str]) -> dict[str, dict]:

    log.info("Running readability check (min score: %s)...", READABILITY_MIN)
    results = {}

    skip = {"FLOORING INSTALLATION PROPOSAL", "Pricing Summary", "Terms and Conditions"}

    for label, text in sections.items():
        if label in skip or len(text.split()) < 10:
            continue

        score = textstat.flesch_reading_ease(text)
        entry = {"score": round(score, 1), "rewrite": None}

        if score < READABILITY_MIN:
            log.info("Section '%s' scored %.1f — requesting LLM rewrite", label, score)
            prompt = REWRITE_PROMPT.format(section_label=label, section_text=text)
            entry["rewrite"] = call_llm(
                prompt, model=T3_LLM_MODEL, base_url=T3_LLM_BASE_URL
            )
        else:
            log.info("Section '%s' scored %.1f — OK", label, score)

        results[label] = entry

    return results


# write reports
def write_reports(
    completeness_issues: list[str],
    consistency_issues:  list[str],
    contradictions:      str,
    readability:         dict[str, dict],
):
    today = date.today().isoformat()
    os.makedirs(os.path.dirname(COMPLETENESS_REPORT), exist_ok=True)

    with open(COMPLETENESS_REPORT, "w") as f:
        f.write(f"COMPLETENESS REPORT — {today}\n{'='*50}\n\n")
        if completeness_issues:
            f.write("\n".join(completeness_issues))
        else:
            f.write("All required sections present and non-empty.")
    log.info("Completeness report → %s", COMPLETENESS_REPORT)

    with open(CONSISTENCY_REPORT, "w") as f:
        f.write(f"CONSISTENCY REPORT — {today}\n{'='*50}\n\n")
        if consistency_issues:
            f.write("\n".join(consistency_issues))
        else:
            f.write("All dollar figures verified against Team 2 cost data.")
    log.info("Consistency report → %s", CONSISTENCY_REPORT)

    with open(READABILITY_REPORT, "w") as f:
        f.write(f"READABILITY REPORT — {today}\n{'='*50}\n\n")
        f.write(f"LLM CONTRADICTION CHECK:\n{contradictions}\n\n")
        f.write(f"{'='*50}\nREADABILITY SCORES (Flesch Reading Ease, min={READABILITY_MIN}):\n\n")
        for label, data in readability.items():
            f.write(f"[{data['score']}] {label}\n")
            if data["rewrite"]:
                f.write(f"  SUGGESTED REWRITE:\n  {data['rewrite']}\n")
            f.write("\n")
    log.info("Readability report → %s", READABILITY_REPORT)


# def save_validated_docx(draft_path: str, readability: dict[str, dict], out_path: str):

#     shutil.copy2(draft_path, out_path)

#     doc = Document(out_path)
#     doc.add_page_break()
#     doc.add_heading("Track 3 Validation Summary", level=1)

#     p = doc.add_paragraph()
#     p.add_run(f"Validated: {date.today().isoformat()}").bold = True

#     doc.add_heading("Readability Scores", level=2)
#     for label, data in readability.items():
#         line = f"{label}: {data['score']}"
#         if data["rewrite"]:
#             line += "  ⚠ Rewrite suggested (see readability_report.txt)"
#         doc.add_paragraph(line, style="List Bullet")

#     doc.save(out_path)
#     log.info("Validated proposal → %s", out_path)


# track 2
def run_track2(
    team1_path: str = "data/dummy_team1_output.json",
    team2_path: str = "data/dummy_team2_output.json",
) -> str:
    
    log.info("=" * 60)
    log.info("TRACK 2  —  CONTENT GENERATION")
    log.info("Model: %s  |  Output: %s", LLM_MODEL, OUTPUT_PATH)
    log.info("=" * 60)

    team1, team2 = load_inputs(team1_path, team2_path)

    ctx = aggregate_scope(team1, team2)

    generated: dict[str, str] = {}
    for section in SCHEMA["sections"]:
        if section["llm"]:
            sid = section["id"]
            if sid == "alternate_pricing" and not ctx["alternates"]:
                log.info("No alternates — skipping alternate_pricing section")
                continue
            generated[sid] = generate_section(sid, ctx)

    doc = assemble_document(ctx, generated)

    save_document(doc, OUTPUT_PATH)

    log.info("✓ Track 2 complete.")
    return OUTPUT_PATH


# track 3
def run_track3(team2_path: str = "data/dummy_team2_output.json") -> str:

    log.info("=" * 60)
    log.info("TRACK 3  —  REVIEW & VALIDATION")
    log.info("Draft: %s", DRAFT_PATH)
    log.info("=" * 60)

    sections = load_draft(DRAFT_PATH)
    team2    = load_team2(team2_path)

    completeness_issues = check_completeness(sections)

    consistency_issues = check_consistency(sections, team2)

    contradictions = check_contradictions(sections)

    readability = check_readability(sections)

    write_reports(completeness_issues, consistency_issues, contradictions, readability)
    # save_validated_docx(DRAFT_PATH, readability, VALIDATED_PATH)

    log.info("✓ Track 3 complete.")
    return READABILITY_REPORT
